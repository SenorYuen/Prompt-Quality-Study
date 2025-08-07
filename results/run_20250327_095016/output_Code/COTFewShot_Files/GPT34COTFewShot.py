from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path

    def read_text(self):
        try:
            doc = Document(self.file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except:
            return None

    def write_text(self, content, font_size=12, alignment='left'):
        try:
            doc = Document()
            paragraph = doc.add_paragraph(content)
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            paragraph.alignment = self._get_alignment_value(alignment)
            doc.save(self.file_path)
            return True
        except:
            return False

    def add_heading(self, heading, level=1):
        try:
            doc = Document(self.file_path)
            doc.add_heading(heading, level=level)
            doc.save(self.file_path)
            return True
        except:
            return False

    def add_table(self, data):
        try:
            doc = Document(self.file_path)
            table = doc.add_table(rows=1, cols=len(data[0]))
            hdr_cells = table.rows[0].cells
            for i, heading in enumerate(data[0]):
                hdr_cells[i].text = heading
            for row_data in data[1:]:
                row_cells = table.add_row().cells
                for i, item in enumerate(row_data):
                    row_cells[i].text = str(item)
            doc.save(self.file_path)
            return True
        except:
            return False

    def _get_alignment_value(self, alignment):
        if alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            return WD_PARAGRAPH_ALIGNMENT.LEFT