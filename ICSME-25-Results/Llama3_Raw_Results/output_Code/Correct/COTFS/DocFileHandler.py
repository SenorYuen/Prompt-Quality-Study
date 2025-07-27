from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path

    def read_text(self):
        try:
            doc = Document(self.file_path)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            return '\n'.join(text)
        except:
            return None

    def write_text(self, content, font_size=12, alignment='left'):
        try:
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run(content)
            font = run.font
            font.size = Pt(font_size)
            para.alignment = self._get_alignment_value(alignment)
            doc.save(self.file_path)
            return True
        except:
            return False

    def add_heading(self, heading, level=1):
        try:
            doc = Document()
            para = doc.add_heading(heading, level=level)
            doc.save(self.file_path)
            return True
        except:
            return False

    def add_table(self, data):
        try:
            doc = Document()
            table = doc.add_table(rows=len(data), cols=len(data[0]), style='Table Grid')
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = str(cell)
            doc.save(self.file_path)
            return True
        except:
            return False

    def _get_alignment_value(self, alignment):
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            return WD_PARAGRAPH_ALIGNMENT.LEFT