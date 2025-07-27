from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocFileHandler:
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.

        Args:
            file_path (str): The path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.

        Returns:
            str: The content of the Word document.
        """
        doc = Document(self.file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return '\n'.join(text)

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.

        Args:
            content (str): The text content to write.
            font_size (int, optional): The font size of the text. Defaults to 12.
            alignment (str, optional): The alignment of the text ('left', 'center', or 'right'). Defaults to 'left'.

        Returns:
            bool: True if the write operation is successful, False otherwise.
        """
        try:
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run(content)
            font = run.font
            font.size = Pt(font_size)
            para.alignment = self._get_alignment_value(alignment)
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error: {e}")
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document with the specified text.

        Args:
            heading (str): The text of the heading.
            level (int, optional): The level of the heading. Defaults to 1.

        Returns:
            bool: True if the heading is successfully added, False otherwise.
        """
        try:
            doc = Document()
            doc.add_heading(heading, level=level)
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error: {e}")
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.

        Args:
            data (list of lists): The data to populate the table.

        Returns:
            bool: True if the table is successfully added, False otherwise.
        """
        try:
            doc = Document()
            table = doc.add_table(rows=len(data), cols=len(data[0]))
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = str(cell)
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error: {e}")
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.

        Args:
            alignment (str): The alignment string ('left', 'center', or 'right').

        Returns:
            int: The alignment value.
        """
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            raise ValueError("Invalid alignment value")