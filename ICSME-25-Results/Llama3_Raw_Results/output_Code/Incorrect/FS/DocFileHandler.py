from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    # Method to read the content of a Word document
    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        try:
            # Open the Word document
            document = Document(self.file_path)
            # Initialize an empty string to store the content
            content = ""
            # Iterate over each paragraph in the document
            for paragraph in document.paragraphs:
                # Append the paragraph text to the content string
                content += paragraph.text + "\n"
            # Return the content string
            return content
        except Exception as e:
            # Handle any exceptions that occur during the read operation
            print(f"Error reading document: {e}")
            return None

    # Method to write text to a Word document
    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            # Create a new Word document or open an existing one
            document = Document(self.file_path)
            # Add a paragraph to the document
            paragraph = document.add_paragraph()
            # Set the paragraph alignment
            paragraph.alignment = self._get_alignment_value(alignment)
            # Add a run to the paragraph
            run = paragraph.add_run(content)
            # Set the font size of the run
            font = run.font
            font.size = Pt(font_size)
            # Save the document
            document.save(self.file_path)
            # Return True to indicate a successful write operation
            return True
        except Exception as e:
            # Handle any exceptions that occur during the write operation
            print(f"Error writing to document: {e}")
            return False

    # Method to add a heading to a Word document
    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            # Create a new Word document or open an existing one
            document = Document(self.file_path)
            # Add a heading to the document
            document.add_heading(heading, level=level)
            # Save the document
            document.save(self.file_path)
            # Return True to indicate a successful add operation
            return True
        except Exception as e:
            # Handle any exceptions that occur during the add operation
            print(f"Error adding heading to document: {e}")
            return False

    # Method to add a table to a Word document
    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            # Create a new Word document or open an existing one
            document = Document(self.file_path)
            # Add a table to the document
            table = document.add_table(rows=len(data), cols=len(data[0]))
            # Populate the table with data
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = str(cell)
            # Save the document
            document.save(self.file_path)
            # Return True to indicate a successful add operation
            return True
        except Exception as e:
            # Handle any exceptions that occur during the add operation
            print(f"Error adding table to document: {e}")
            return False

    # Helper method to get the alignment value from a string
    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            return WD_PARAGRAPH_ALIGNMENT.LEFT