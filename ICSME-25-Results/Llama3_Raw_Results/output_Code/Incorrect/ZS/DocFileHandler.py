# This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    def __init__(self, file_path):
        # Initializes the DocFileHandler object with the specified file path.
        self.file_path = file_path
        self.document = Document(self.file_path)

    def read_text(self):
        # Reads the content of a Word document and returns it as a string.
        text = ''
        for paragraph in self.document.paragraphs:
            text += paragraph.text + '\n'
        return text

    def write_text(self, content, font_size=12, alignment='left'):
        # Writes the specified content to a Word document.
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(content)
        font = run.font
        font.size = Pt(font_size)
        paragraph.alignment = self._get_alignment_value(alignment)
        self.document.save(self.file_path)
        return True

    def add_heading(self, heading, level=1):
        # Adds a heading to the Word document.
        self.document.add_heading(heading, level=level)
        self.document.save(self.file_path)
        return True

    def add_table(self, data):
        # Adds a table to the Word document with the specified data.
        table = self.document.add_table(rows=len(data), cols=len(data[0]), style='Table Grid')
        for i, row in enumerate(data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)
        self.document.save(self.file_path)
        return True

    def _get_alignment_value(self, alignment):
        # Returns the alignment value corresponding to the given alignment string.
        alignment_values = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        return alignment_values.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)


# Example usage:
if __name__ == "__main__":
    file_path = "example.docx"
    handler = DocFileHandler(file_path)
    print(handler.read_text())
    handler.write_text("Hello, World!", font_size=14, alignment='center')
    handler.add_heading("Example Heading", level=2)
    data = [[1, 2, 3], [4, 5, 6], [7, 8, 9]]
    handler.add_table(data)