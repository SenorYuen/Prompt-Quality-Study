from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path

    def read_text(self):
        document = Document(self.file_path)
        full_text = []
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)

    def write_text(self, content, font_size=12, alignment='left'):
        try:
            document = Document()
            paragraph = document.add_paragraph(content)
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            paragraph.alignment = self._get_alignment_value(alignment)
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error writing text: {e}")
            return False

    def add_heading(self, heading, level=1):
        try:
            document = Document(self.file_path)
            document.add_heading(heading, level=level)
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False

    def add_table(self, data):
        try:
            document = Document(self.file_path)
            table = document.add_table(rows=1, cols=len(data[0]))
            hdr_cells = table.rows[0].cells
            for i, heading in enumerate(data[0]):
                hdr_cells[i].text = heading
            for row_data in data[1:]:
                row_cells = table.add_row().cells
                for i, item in enumerate(row_data):
                    row_cells[i].text = item
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False

    def _get_alignment_value(self, alignment):
        if alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            return WD_PARAGRAPH_ALIGNMENT.LEFT