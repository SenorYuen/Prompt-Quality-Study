from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocFileHandler:
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        """
        self.file_path = file_path
        self.document = Document()  # Create a new Document

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        doc = Document(self.file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            paragraph = self.document.add_paragraph(content)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = self._get_alignment_value(alignment)
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            self.document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error writing text: {e}")
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            self.document.add_heading(heading, level=level)
            self.document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            table = self.document.add_table(rows=1, cols=len(data[0]))
            hdr_cells = table.rows[0].cells
            for i, heading in enumerate(data[0]):
                hdr_cells[i].text = heading

            for row_data in data[1:]:
                row_cells = table.add_row().cells
                for i, item in enumerate(row_data):
                    row_cells[i].text = str(item)

            self.document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :return: int, the alignment value.
        """
        alignment_dict = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        return alignment_dict.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)